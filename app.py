import os
import subprocess
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

# Tải biến môi trường từ file .env
load_dotenv()

# Cấu hình client OpenAI
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")  # Lấy API Key từ biến môi trường
)

# Cấu hình ứng dụng Flask
app = Flask(__name__, template_folder='templates')  # Đảm bảo đường dẫn tới thư mục templates

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {
    'wav', 'mp3', 'flac', 'm4a', 'aac', 'mp4', 'mov', 'avi', 'mkv', 'amsr', 'ogg'
}

# Đảm bảo thư mục uploads tồn tại
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Cấu hình email server
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USERNAME = os.getenv("SMTP_USERNAME")  # Email của bạn
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")  # Mật khẩu email của bạn

# Kiểm tra định dạng file hợp lệ
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Chuyển đổi định dạng âm thanh/video sang WAV bằng FFmpeg
def convert_audio(input_path, output_path):
    try:
        processed_output_path = f"{os.path.splitext(output_path)[0]}_processed.wav"
        subprocess.run([ 
            'ffmpeg', '-y', '-i', input_path, 
            '-ar', '16000',  # Tần số lấy mẫu 
            '-ac', '1',      # Mono channel 
            '-f', 'wav', 
            processed_output_path 
        ], check=True)
        return processed_output_path
    except subprocess.CalledProcessError as e:
        print(f"FFmpeg error: {e}")
        return None

# Chuyển đổi giọng nói thành văn bản bằng Whisper API, bao gồm timestamp
def transcribe_audio(file_path):
    try:
        with open(file_path, 'rb') as audio_file:
            response = client.audio.transcriptions.create(
                model="whisper-1",  # Bạn có thể thay đổi model nếu cần
                file=audio_file,
                language="vi"  # Ngôn ngữ: tiếng Việt
            )
        return response.text  # Trả về văn bản hoàn chỉnh
    except Exception as e:
        print(f"Whisper API error for file {file_path}: {e}")
        return "..."

# Dịch song ngữ bằng ChatGPT
def translate_bilingual(text):
    if not text or text.strip() == "...":
        return "..."
    
    prompt = f"""
    Bạn nhận được đầu ra từ Whisper API. Nếu có phần không rõ ràng, hãy thay thế bằng dấu "...". Nhiệm vụ của bạn là:

    1. Chỉnh sửa nội dung để rõ ràng, dễ đọc hơn.
    2. Cắt nhỏ, mỗi đoạn không quá 20 từ.
    3. Dịch nội dung sang tiếng Anh theo định dạng:
       - : <Câu gốc tiếng Việt, nếu không rõ thì giữ nguyên dấu "...">
       -> B : <Bản dịch tiếng Anh, nếu không rõ thì giữ nguyên "...">
    Nội dung: "{text}"
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "system", "content": "Bạn là chuyên gia dịch thuật..."},
                      {"role": "user", "content": prompt}],
            max_tokens=3000,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"ChatGPT API error: {e}")
        return "..."
# Xuất kết quả ra file Word (không có timestamp)
def export_to_word(bilingual_text, output_file):
    try:
        document = Document()
        document.add_heading("Bản Dịch Song Ngữ", level=1)
        document.add_paragraph(bilingual_text)
        document.save(output_file)
        return output_file
    except Exception as e:
        print(f"Lỗi khi xuất file Word: {e}")
        return None

# Gửi email với file đính kèm
def send_email_with_attachment(to_email, subject, body, attachment_paths):
    try:
        msg = MIMEMultipart()
        msg['From'] = SMTP_USERNAME
        msg['To'] = to_email
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        for attachment_path in attachment_paths:
            with open(attachment_path, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename={os.path.basename(attachment_path)}'
                )
                msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USERNAME, SMTP_PASSWORD)
            server.send_message(msg)
        print(f"Email đã được gửi tới {to_email}")
    except Exception as e:
        print(f"Không thể gửi email: {e}")

# Trang chính của ứng dụng
@app.route('/')
def index():
    return render_template('index.html')

# Xử lý upload file
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'audio' not in request.files or 'email' not in request.form:
        return jsonify({'error': 'Không tìm thấy file âm thanh hoặc email.'}), 400

    file = request.files['audio']
    email = request.form['email']
    if file.filename == '' or not email:
        return jsonify({'error': 'Không có file nào được chọn hoặc email không hợp lệ.'}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        output_filename = f"{os.path.splitext(filename)[0]}.wav"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        converted = convert_audio(input_path, output_path)
        if not converted:
            return jsonify({'error': 'Lỗi trong quá trình chuyển đổi định dạng âm thanh/video.'}), 500

        processed_output_path = converted

        transcript = transcribe_audio(processed_output_path)

        if not transcript:
            return jsonify({'error': 'Lỗi trong quá trình phiên âm âm thanh.'}), 500

        bilingual_text = translate_bilingual(transcript)
        if not bilingual_text:
            return jsonify({'error': 'Lỗi trong quá trình dịch song ngữ.'}), 500

        word_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}_result.docx")
        export_to_word(bilingual_text, word_file_path)

        send_email_with_attachment(
            to_email=email,
            subject="Kết quả dịch âm thanh",
            body="Vui lòng kiểm tra file đính kèm để xem kết quả.",
            attachment_paths=[word_file_path]
        )

        try:
            os.remove(input_path)
            os.remove(processed_output_path)
        except Exception as e:
            print(f"Error deleting files: {e}")

        return jsonify({'message': 'Kết quả đã được gửi qua email.', 'bilingual_text': bilingual_text})

    else:
        return jsonify({'error': 'Định dạng file không được hỗ trợ.'}), 400

if __name__ == '__main__':
    app.run(debug=True)
