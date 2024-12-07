import os
import subprocess
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from flask import Flask, request, jsonify, render_template
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from openai import OpenAI
from docx import Document

# Tải biến môi trường từ file .env
load_dotenv()

# Cấu hình client OpenAI
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")  # Lấy API Key từ biến môi trường
)

# Cấu hình ứng dụng Flask
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {
    'wav', 'mp3', 'flac', 'm4a', 'aac', 'mp4', 'mov', 'avi', 'mkv', 'amsr', 'ogg'
}

# Đảm bảo thư mục uploads tồn tại
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Cấu hình email server
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USERNAME = os.getenv("SMTP_USERNAME")  # Email của bạn
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD")  # Mật khẩu email của bạn

# Kiểm tra định dạng file hợp lệ
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# Chuyển đổi định dạng âm thanh/video sang WAV bằng FFmpeg
def convert_audio(input_path, output_path):
    try:
        processed_output_path = f"{os.path.splitext(output_path)[0]}_processed.wav"
        subprocess.run([
            'ffmpeg', '-y', '-i', input_path,
            '-ar', '16000',  # Tần số lấy mẫu
            '-ac', '1',       # Mono channel
            '-f', 'wav',
            processed_output_path
        ], check=True)
        if os.path.exists(processed_output_path):
            print(f"Chuyển đổi thành công: {processed_output_path}")
            return processed_output_path
        else:
            print(f"Chuyển đổi thất bại: {processed_output_path} không tồn tại.")
            return None
    except subprocess.CalledProcessError as e:
        print(f"FFmpeg error: {e}")
        return None

# Chuyển đổi giọng nói thành văn bản bằng Whisper API
def transcribe_audio(file_path):
    try:
        with open(file_path, 'rb') as audio_file:
            response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language="vi",  # Ngôn ngữ: tiếng Việt
                response_format="verbose_json"  # Định dạng phản hồi bao gồm timestamps
            )
        
        # Kiểm tra xem phản hồi có chứa thuộc tính 'segments' không
        if hasattr(response, 'segments') and response.segments:
            transcript_with_timestamps = []
            for segment in response.segments:
                start_time = segment.start
                end_time = segment.end
                text = segment.text.strip()
                transcript_with_timestamps.append({
                    "start": start_time,
                    "end": end_time,
                    "text": text
                })
            return transcript_with_timestamps
        else:
            # Nếu không có segments, chỉ lấy toàn bộ văn bản
            return [{"start": 0, "end": 0, "text": response.text.strip()}]
    except Exception as e:
        print(f"Whisper API error for file {file_path}: {e}")
        return []

# Dịch song ngữ bằng ChatGPT
def translate_bilingual(transcript_with_timestamps):
    if not transcript_with_timestamps:
        return "..."

    bilingual_segments = []
    
    for segment in transcript_with_timestamps:
        vietnamese_text = segment['text']
        end_time = segment['end']
        
        # Chuyển end_time thành định dạng phút:giây
        end_min = int(end_time // 60)
        end_sec = int(end_time % 60)
        time_formatted = f"[{end_min}:{end_sec:02}]"
        
        prompt = f"""
Bạn nhận được đoạn văn bản sau đây từ Whisper API. Nếu có phần không rõ ràng, hãy thay thế bằng dấu "...".
Nhiệm vụ của bạn là:

1. Chỉnh sửa nội dung để rõ ràng, dễ đọc hơn.
2. Dịch nội dung sang tiếng Anh theo định dạng:
   - {time_formatted}: {vietnamese_text}
   -> {time_formatted}: <Bản dịch tiếng Anh>

   Làm theo mẫu sau:
-  [0:02]: Anh nhớ, đừng lấy những cái nót ở đó.
-> [0:02]: Remember, don't take those notes there.

-  [0:04]: Anh nhớ, đừng lấy những cái nót ở đó.
-> [0:04]: Remember, don't take those notes there.
Nội dung: "{vietnamese_text}"
"""
        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Bạn là chuyên gia dịch thuật."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=3000,
                temperature=0.3,
            )
            # Đảm bảo truy cập đúng cấu trúc phản hồi
            translated_text = response.choices[0].message.content.strip()
            bilingual_segments.append(translated_text)
        except Exception as e:
            print(f"ChatGPT API error: {e}")
            bilingual_segments.append(f"A (Tiếng Việt) {time_formatted}: {vietnamese_text}\nB (English) {time_formatted}: ...")
    
    return "\n".join(bilingual_segments)


# Xuất kết quả ra file Word
def export_to_word(bilingual_text, output_file):
    try:
        document = Document()
        document.add_heading("Bản Dịch Song Ngữ", level=1)
        
        bilingual_lines = bilingual_text.split('\n')
        for line in bilingual_lines:
            if line.strip():  # Kiểm tra nếu dòng không rỗng
                document.add_paragraph(line)
        
        document.save(output_file)
        print(f"Xuất file Word thành công: {output_file}")
        return output_file
    except Exception as e:
        print(f"Lỗi khi xuất file Word: {e}")
        return None

# Gửi email với file đính kèm
def send_email_with_attachment(to_email, subject, body, attachment_paths):
    try:
        msg = MIMEMultipart()
        msg['From'] = SMTP_USERNAME
        msg['To'] = to_email
        msg['Subject'] = subject

        msg.attach(MIMEText(body, 'plain'))

        for attachment_path in attachment_paths:
            if not os.path.isfile(attachment_path):
                print(f"File không tồn tại để đính kèm: {attachment_path}")
                continue
            with open(attachment_path, 'rb') as attachment:
                part = MIMEBase('application', 'octet-stream')
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                part.add_header(
                    'Content-Disposition',
                    f'attachment; filename={os.path.basename(attachment_path)}'
                )
                msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USERNAME, SMTP_PASSWORD)
            server.send_message(msg)
        print(f"Email đã được gửi tới {to_email}")
    except Exception as e:
        print(f"Không thể gửi email: {e}")

# Trang chính của ứng dụng
@app.route('/')
def index():
    return render_template('index.html')

# Xử lý upload file
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'audio' not in request.files or 'email' not in request.form:
        return jsonify({'error': 'Không tìm thấy file âm thanh hoặc email.'}), 400

    files = request.files.getlist('audio')
    email = request.form['email']
    if not files or not email:
        return jsonify({'error': 'Không có file nào được chọn hoặc email không hợp lệ.'}), 400

    all_results = {}
    attachments = []

    for idx, file in enumerate(files):
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)
            print(f"Đã tải lên file: {input_path}")

            output_filename = f"{os.path.splitext(filename)[0]}.wav"
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
            converted = convert_audio(input_path, output_path)
            if not converted:
                return jsonify({'error': f'Lỗi trong quá trình chuyển đổi định dạng âm thanh/video của file {filename}.'}), 500

            # Chuyển đổi WAV thành văn bản
            transcript_with_timestamps = transcribe_audio(converted)
            if not transcript_with_timestamps:
                return jsonify({'error': f'Lỗi trong quá trình phiên âm âm thanh của file {filename}.'}), 500

            bilingual_text = translate_bilingual(transcript_with_timestamps)
            if not bilingual_text:
                return jsonify({'error': f'Lỗi trong quá trình dịch song ngữ của file {filename}.'}), 500

            word_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}_result.docx")
            export_status = export_to_word(bilingual_text, word_file_path)
            if not export_status:
                return jsonify({'error': f'Lỗi trong quá trình xuất file Word của file {filename}.'}), 500

            attachments.append(word_file_path)
            print(f"Đã thêm file Word vào đính kèm: {word_file_path}")

            # Lưu kết quả vào biến tạm
            variable_name = f"file_{idx + 1}"
            all_results[variable_name] = {
                "bilingual_text": bilingual_text
            }

            # Xóa file gốc và file đã chuyển đổi
            try:
                if os.path.exists(input_path):
                    os.remove(input_path)
                    print(f"Đã xóa file gốc: {input_path}")
                if os.path.exists(converted):
                    os.remove(converted)
                    print(f"Đã xóa file chuyển đổi: {converted}")
            except Exception as e:
                print(f"Error deleting files: {e}")

    if not attachments:
        return jsonify({'error': 'Không có file nào được xử lý thành công để gửi email.'}), 500

    # Gửi email với các file Word đính kèm
    send_email_with_attachment(
        to_email=email,
        subject="Kết quả dịch âm thanh",
        body="Vui lòng kiểm tra file đính kèm để xem kết quả.",
        attachment_paths=attachments
    )

    return jsonify({'message': 'Kết quả đã được gửi qua email.', 'results': all_results})

if __name__ == '__main__':
    app.run(debug=True)
